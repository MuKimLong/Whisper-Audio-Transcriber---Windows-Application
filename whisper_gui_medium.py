import tkinter as tk
from tkinter import filedialog, ttk, messagebox
import threading
import whisper
from docx import Document
import os
import torch
import subprocess

def is_gpu_available():
    return torch.cuda.is_available()

# Modeli sadece bir kez yükle
model_size = "medium"  # Hızı artırmak için "small" veya "medium" seçilebilir
global_model = whisper.load_model(model_size)
if is_gpu_available():
    global_model = global_model.to("cuda").half()  # GPU hızlandırması

def convert_audio(input_path):
    output_path = os.path.splitext(input_path)[0] + "_converted.wav"
    command = ["ffmpeg", "-i", input_path, "-ar", "16000", "-ac", "1", output_path, "-y"]
    subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    return output_path

def transcribe_audio():
    file_path = filedialog.askopenfilename(filetypes=[("Audio Files", "*.mp3 *.wav *.m4a")])
    if not file_path:
        return
    
    progress_bar.start()
    transcribe_thread = threading.Thread(target=process_transcription, args=(file_path,))
    transcribe_thread.start()

def process_transcription(file_path):
    try:
        converted_path = convert_audio(file_path)
        result = global_model.transcribe(converted_path)
        text = result["text"]
        save_to_word(file_path, text)
        messagebox.showinfo("Başarılı", "Deşifre tamamlandı ve kaydedildi!")
    except Exception as e:
        messagebox.showerror("Hata", f"Bir hata oluştu: {e}")
    finally:
        progress_bar.stop()

def save_to_word(file_path, text):
    file_name = os.path.splitext(os.path.basename(file_path))[0]
    doc = Document()
    doc.add_heading(file_name, level=1)
    
    sentences = text.split(". ")
    for sentence in sentences:
        doc.add_paragraph(sentence)
    
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], title="Dosyayı Kaydet")
    if save_path:
        doc.save(save_path)

def create_gui():
    global progress_bar
    
    root = tk.Tk()
    root.title("Whisper Transcribe")
    root.geometry("400x200")
    
    label = tk.Label(root, text="Ses dosyanızı seçin ve deşifre edin:")
    label.pack(pady=10)
    
    select_button = tk.Button(root, text="Dosya Seç", command=transcribe_audio)
    select_button.pack(pady=5)
    
    progress_bar = ttk.Progressbar(root, mode="indeterminate")
    progress_bar.pack(pady=10, fill=tk.X, padx=20)
    
    root.mainloop()

if __name__ == "__main__":
    create_gui()